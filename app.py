import os
import sqlite3
import uuid
from datetime import datetime

from flask import (
    Flask,
    request,
    jsonify,
    send_from_directory,
    send_file
)

from werkzeug.utils import secure_filename
from openpyxl import load_workbook, Workbook
from docx import Document


BASE = os.path.dirname(
    os.path.abspath(__file__)
)

DB = os.path.join(
    BASE,
    "employees.db"
)

UPLOADS = os.path.join(
    BASE,
    "uploads"
)

OUTPUTS = os.path.join(
    BASE,
    "documents"
)

os.makedirs(UPLOADS, exist_ok=True)
os.makedirs(OUTPUTS, exist_ok=True)


app = Flask(__name__)

app.config[
    "MAX_CONTENT_LENGTH"
] = 100 * 1024 * 1024


def db():

    c = sqlite3.connect(DB)
    c.row_factory = sqlite3.Row
    return c


def init_db():

    c = db()

    c.execute("""
    CREATE TABLE IF NOT EXISTS employees(

        id INTEGER PRIMARY KEY AUTOINCREMENT,

        card_number TEXT,
        first_name TEXT,
        last_name TEXT,
        rank TEXT,

        installation_date TEXT,

        workplace TEXT,
        current_location TEXT,

        birth_date TEXT,
        phone TEXT,
        email TEXT,

        note TEXT,

        photo TEXT,
        employee_file TEXT,
        note_file TEXT,

        created_at TEXT
    )
    """)

    c.commit()
    c.close()


init_db()


@app.route("/")
def index():

    return send_from_directory(
        BASE,
        "index.html"
    )


@app.route("/uploads/<path:name>")
def upload_file(name):

    return send_from_directory(
        UPLOADS,
        name
    )


@app.route("/documents/<path:name>")
def document_file(name):

    return send_from_directory(
        OUTPUTS,
        name
    )


def save_upload(file):

    if not file:
        return ""

    if not file.filename:
        return ""

    original = secure_filename(
        file.filename
    )

    ext = ""

    if "." in original:
        ext = "." + original.rsplit(
            ".",
            1
        )[1].lower()

    filename = (
        uuid.uuid4().hex
        + ext
    )

    path = os.path.join(
        UPLOADS,
        filename
    )

    file.save(path)

    return filename


def public_file(name):

    if not name:
        return ""

    return "/uploads/" + name


@app.get("/api/employees")
def employees():

    q = request.args.get(
        "q",
        ""
    ).strip()

    rank = request.args.get(
        "rank",
        ""
    ).strip()

    place = request.args.get(
        "place",
        ""
    ).strip()

    sort = request.args.get(
        "sort",
        ""
    )

    c = db()

    sql = """
    SELECT *
    FROM employees
    WHERE 1=1
    """

    values = []

    if q:

        sql += """
        AND (
            first_name LIKE ?
            OR last_name LIKE ?
            OR card_number LIKE ?
            OR rank LIKE ?
            OR workplace LIKE ?
            OR current_location LIKE ?
            OR phone LIKE ?
        )
        """

        x = "%" + q + "%"

        values += [
            x,x,x,x,x,x,x
        ]

    if rank:

        sql += """
        AND rank LIKE ?
        """

        values.append(
            "%" + rank + "%"
        )

    if place:

        sql += """
        AND (
            workplace LIKE ?
            OR current_location LIKE ?
        )
        """

        values += [
            "%" + place + "%",
            "%" + place + "%"
        ]

    order = {
        "name":
            "first_name COLLATE NOCASE",

        "rank":
            "rank COLLATE NOCASE",

        "place":
            "workplace COLLATE NOCASE"
    }.get(
        sort,
        "created_at DESC"
    )

    sql += " ORDER BY " + order

    rows = c.execute(
        sql,
        values
    ).fetchall()

    total = c.execute(
        "SELECT COUNT(*) FROM employees"
    ).fetchone()[0]

    files = c.execute("""
        SELECT COUNT(*)
        FROM employees
        WHERE employee_file IS NOT NULL
        AND employee_file != ''
    """).fetchone()[0]

    photos = c.execute("""
        SELECT COUNT(*)
        FROM employees
        WHERE photo IS NOT NULL
        AND photo != ''
    """).fetchone()[0]

    notes = c.execute("""
        SELECT COUNT(*)
        FROM employees
        WHERE note IS NOT NULL
        AND note != ''
    """).fetchone()[0]

    c.close()

    result = []

    for r in rows:

        x = dict(r)

        x["photo"] = public_file(
            x["photo"]
        )

        x["employee_file"] = public_file(
            x["employee_file"]
        )

        x["note_file"] = public_file(
            x["note_file"]
        )

        result.append(x)

    return jsonify({
        "items": result,
        "stats": {
            "employees": total,
            "files": files,
            "photos": photos,
            "notes": notes,
            "date": datetime.now().strftime(
                "%Y-%m-%d"
            )
        }
    })


@app.post("/api/employees")
def create_employee():

    form = request.form

    photo = save_upload(
        request.files.get("photo")
    )

    employee_file = save_upload(
        request.files.get("employee_file")
    )

    note_file = save_upload(
        request.files.get("note_file")
    )

    c = db()

    cur = c.execute("""
    INSERT INTO employees(
        card_number,
        first_name,
        last_name,
        rank,
        installation_date,
        workplace,
        current_location,
        birth_date,
        phone,
        email,
        note,
        photo,
        employee_file,
        note_file,
        created_at
    )
    VALUES(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
    """, (

        form.get("card_number",""),
        form.get("first_name",""),
        form.get("last_name",""),
        form.get("rank",""),

        form.get(
            "installation_date",
            ""
        ),

        form.get(
            "workplace",
            ""
        ),

        form.get(
            "current_location",
            ""
        ),

        form.get(
            "birth_date",
            ""
        ),

        form.get(
            "phone",
            ""
        ),

        form.get(
            "email",
            ""
        ),

        form.get(
            "note",
            ""
        ),

        photo,
        employee_file,
        note_file,

        datetime.now().isoformat()
    ))

    c.commit()

    employee_id = cur.lastrowid

    c.close()

    return jsonify({
        "ok": True,
        "id": employee_id
    })


@app.put("/api/employees/<int:employee_id>")
def update_employee(employee_id):

    c = db()

    old = c.execute(
        "SELECT * FROM employees WHERE id=?",
        (employee_id,)
    ).fetchone()

    if not old:

        c.close()

        return jsonify({
            "error":
                "الموظف غير موجود"
        }), 404

    form = request.form

    photo = old["photo"]

    employee_file = old[
        "employee_file"
    ]

    note_file = old[
        "note_file"
    ]

    new_photo = request.files.get(
        "photo"
    )

    new_employee_file = request.files.get(
        "employee_file"
    )

    new_note_file = request.files.get(
        "note_file"
    )

    if new_photo and new_photo.filename:
        photo = save_upload(
            new_photo
        )

    if new_employee_file and new_employee_file.filename:
        employee_file = save_upload(
            new_employee_file
        )

    if new_note_file and new_note_file.filename:
        note_file = save_upload(
            new_note_file
        )

    c.execute("""
    UPDATE employees SET

        card_number=?,
        first_name=?,
        last_name=?,
        rank=?,
        installation_date=?,
        workplace=?,
        current_location=?,
        birth_date=?,
        phone=?,
        email=?,
        note=?,
        photo=?,
        employee_file=?,
        note_file=?

    WHERE id=?
    """, (

        form.get("card_number",""),
        form.get("first_name",""),
        form.get("last_name",""),
        form.get("rank",""),
        form.get(
            "installation_date",
            ""
        ),
        form.get(
            "workplace",
            ""
        ),
        form.get(
            "current_location",
            ""
        ),
        form.get(
            "birth_date",
            ""
        ),
        form.get(
            "phone",
            ""
        ),
        form.get(
            "email",
            ""
        ),
        form.get(
            "note",
            ""
        ),
        photo,
        employee_file,
        note_file,

        employee_id
    ))

    c.commit()
    c.close()

    return jsonify({
        "ok": True
    })


@app.delete("/api/employees/<int:employee_id>")
def delete_employee(employee_id):

    c = db()

    c.execute(
        "DELETE FROM employees WHERE id=?",
        (employee_id,)
    )

    c.commit()
    c.close()

    return jsonify({
        "ok": True
    })


# ------------------------------------------------
# قراءة Excel
# ------------------------------------------------

@app.post("/api/excel/read")
def excel_read():

    file = request.files.get(
        "excel"
    )

    if not file or not file.filename:

        return jsonify({
            "error":
                "لم يتم اختيار ملف Excel"
        }), 400

    filename = secure_filename(
        file.filename
    )

    path = os.path.join(
        UPLOADS,
        uuid.uuid4().hex
        + "_"
        + filename
    )

    file.save(path)

    sheet = request.form.get(
        "sheet",
        "SS"
    )

    wilaya = request.form.get(
        "wilaya",
        ""
    )

    try:

        wb = load_workbook(
            path,
            read_only=True,
            data_only=True
        )

        if sheet not in wb.sheetnames:

            return jsonify({
                "error":
                    "صفحة "
                    + sheet
                    + " غير موجودة في الملف"
            }), 400

        ws = wb[sheet]

        rows = 0

        for row in ws.iter_rows(
            values_only=True
        ):

            if any(
                x is not None
                for x in row
            ):
                rows += 1

        wb.close()

        return jsonify({
            "ok": True,
            "filename": filename,
            "wilaya": wilaya,
            "sheet": sheet,
            "rows": rows
        })

    except Exception as e:

        return jsonify({
            "error":
                "تعذر قراءة Excel: "
                + str(e)
        }), 400


# ------------------------------------------------
# إنشاء الوثائق
# ------------------------------------------------

def employee_data(employee_id):

    c = db()

    e = c.execute(
        "SELECT * FROM employees WHERE id=?",
        (employee_id,)
    ).fetchone()

    c.close()

    return e


@app.post("/api/documents/create")
def create_document():

    employee_id = request.form.get(
        "employee_id"
    )

    doc_type = request.form.get(
        "document_type",
        "work"
    )

    wilaya = request.form.get(
        "wilaya",
        ""
    )

    sheet = request.form.get(
        "sheet",
        "SS"
    )

    page = request.form.get(
        "page",
        ""
    )

    e = employee_data(
        employee_id
    )

    if not e:

        return jsonify({
            "error":
                "الموظف غير موجود"
        }), 404

    if doc_type == "absence":

        title = "شهادة غياب"

        body = f"""
تشهد المصلحة بأن السيد/السيدة:

{e["first_name"]} {e["last_name"]}

الرتبة: {e["rank"] or "—"}
رقم البطاقة: {e["card_number"] or "—"}
مكان العمل: {e["workplace"] or "—"}

قد تم إعداد هذه الوثيقة بناءً على المعلومات
المسجلة في ملف الموظف.

الولاية: {wilaya or "—"}
المصدر: {sheet}
الصفحة/السطر: {page or "—"}
"""

    elif doc_type == "attendance":

        title = "وثيقة حضور"

        body = f"""
تشهد المصلحة بأن:

{e["first_name"]} {e["last_name"]}

الرتبة: {e["rank"] or "—"}
مكان العمل: {e["workplace"] or "—"}

مسجل ضمن موظفي المصلحة.

الولاية: {wilaya or "—"}
المصدر: {sheet}
الصفحة/السطر: {page or "—"}
"""

    else:

        title = "شهادة عمل"

        body = f"""
تشهد المصلحة بأن السيد/السيدة:

{e["first_name"]} {e["last_name"]}

الرتبة: {e["rank"] or "—"}

يعمل/تعمل بالمصلحة في:

{e["workplace"] or "—"}

وذلك حسب المعلومات الإدارية المسجلة.

الولاية: {wilaya or "—"}
المصدر: {sheet}
الصفحة/السطر: {page or "—"}
"""

    doc = Document()

    section = doc.sections[0]

    section.top_margin = 700000
    section.bottom_margin = 700000
    section.left_margin = 900000
    section.right_margin = 900000

    p = doc.add_paragraph()

    p.alignment = 1

    run = p.add_run(
        title
    )

    run.bold = True
    run.font.size = None

    doc.add_paragraph("")

    for line in body.strip().split("\n"):

        p = doc.add_paragraph()

        p.alignment = 2

        p.add_run(
            line
        )

    doc.add_paragraph("")

    p = doc.add_paragraph()

    p.alignment = 2

    p.add_run(
        "حرر بتاريخ: "
        + datetime.now().strftime(
            "%Y/%m/%d"
        )
    )

    filename = (
        uuid.uuid4().hex
        + ".docx"
    )

    path = os.path.join(
        OUTPUTS,
        filename
    )

    doc.save(path)

    return jsonify({
        "ok": True,
        "file":
            "/documents/"
            + filename
    })


# ------------------------------------------------
# استخراج قائمة الموظفين Excel
# ------------------------------------------------

@app.get("/api/employees/export")
def export_employees():

    c = db()

    rows = c.execute("""
    SELECT
        card_number,
        first_name,
        last_name,
        rank,
        installation_date,
        workplace,
        current_location,
        birth_date,
        phone,
        email,
        note
    FROM employees
    ORDER BY last_name
    """).fetchall()

    c.close()

    wb = Workbook()

    ws = wb.active

    ws.title = "SS"

    headers = [
        "رقم البطاقة",
        "الاسم",
        "اللقب",
        "الرتبة",
        "تاريخ التنصيب",
        "مكان العمل",
        "التواجد الحالي",
        "تاريخ الميلاد",
        "الهاتف",
        "البريد",
        "ملاحظة"
    ]

    ws.append(headers)

    for r in rows:
        ws.append([
            r[x] or ""
            for x in [
                "card_number",
                "first_name",
                "last_name",
                "rank",
                "installation_date",
                "workplace",
                "current_location",
                "birth_date",
                "phone",
                "email",
                "note"
            ]
        ])

    filename = (
        "قائمة_الموظفين_"
        + datetime.now().strftime(
            "%Y%m%d_%H%M%S"
        )
        + ".xlsx"
    )

    path = os.path.join(
        OUTPUTS,
        filename
    )

    wb.save(path)

    return send_file(
        path,
        as_attachment=True,
        download_name=filename
    )


@app.get("/api/health")
def health():

    return jsonify({
        "status":
            "online",
        "database":
            True,
        "excel":
            True,
        "documents":
            True
    })


if __name__ == "__main__":

    port = int(
        os.getenv(
            "PORT",
            "5000"
        )
    )

    app.run(
        host="0.0.0.0",
        port=port,
        debug=False
    )